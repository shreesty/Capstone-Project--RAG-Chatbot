from __future__ import annotations

import argparse
import json
import logging
import os
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Iterable, Optional

import pdfplumber
import pytesseract
import trafilatura
from bs4 import BeautifulSoup
from docx import Document
from langdetect import DetectorFactory, LangDetectException, detect
from PIL import Image, ImageOps
from tqdm import tqdm

# langdetect can be non-deterministic unless we set the seed.
DetectorFactory.seed = 0

SUPPORTED_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp"}
SUPPORTED_TEXT_EXTS = {".txt"}
SUPPORTED_HTML_EXTS = {".html", ".htm"}
SUPPORTED_DOC_EXTS = {".docx"}
SUPPORTED_PDF_EXTS = {".pdf"}

TESSERACT_CMD = os.getenv("TESSERACT_CMD")
if TESSERACT_CMD:
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD


@dataclass
class ExtractedDocument:
    source_path: str
    source_domain: str
    source_url: Optional[str]
    kind: str
    language: str
    text: str
    length: int


def clean_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    cleaned = "\n".join(line for line in lines if line)
    return cleaned.strip()


def normalize_path(raw: str, root: Path) -> Path:
    raw = raw.replace("\\", os.sep)
    path = Path(raw)
    return path if path.is_absolute() else root / path


def detect_language(text: str) -> str:
    try:
        return detect(text)
    except LangDetectException:
        return "unknown"


def extract_text_file(path: Path) -> str:
    return clean_text(path.read_text(encoding="utf-8", errors="ignore"))


def extract_html_file(path: Path) -> str:
    html = path.read_text(encoding="utf-8", errors="ignore")
    extracted = trafilatura.extract(html, include_comments=False, include_tables=False) or ""
    if extracted:
        return clean_text(extracted)

    soup = BeautifulSoup(html, "lxml")
    text = soup.get_text(separator="\n")
    return clean_text(text)


def extract_docx_file(path: Path) -> str:
    document = Document(path)
    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    return clean_text("\n".join(parts))


def ocr_image(img: Image.Image, psm: int = 6) -> str:
    """Run Tesseract OCR with a layout-friendly default page segmentation mode."""
    return pytesseract.image_to_string(img, config=f"--oem 3 --psm {psm}")


def extract_image_file(path: Path) -> str:
    with Image.open(path) as img:
        normalized = ImageOps.exif_transpose(img).convert("RGB")
        gray = normalized.convert("L")
        return clean_text(ocr_image(gray, psm=6))


def extract_pdf_file(path: Path, ocr_page_limit: int = 2) -> str:
    text_parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    full_text = clean_text("\n".join(text_parts))
    if full_text:
        return full_text

    ocr_parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages[:ocr_page_limit]:
            page_image = page.to_image(resolution=300)
            gray = page_image.original.convert("L")
            ocr_parts.append(ocr_image(gray, psm=6))
    return clean_text("\n".join(ocr_parts))


def load_metadata_entries(root: Path) -> Iterable[tuple[Path, str, Optional[str], str]]:
    """Yield (path, domain, url, kind) discovered from metadata.jsonl files."""
    metadata_files = sorted(root.glob("*/metadata.jsonl"))
    for metadata_path in metadata_files:
        domain = metadata_path.parent.name
        with metadata_path.open(encoding="utf-8") as f:
            for line in f:
                try:
                    entry = json.loads(line)
                except json.JSONDecodeError:
                    logging.warning("Skipping bad metadata line in %s", metadata_path)
                    continue

                for key, kind in (("text_path", "html_text"), ("html_path", "html_raw")):
                    raw_path = entry.get(key)
                    if not raw_path:
                        continue
                    path = normalize_path(raw_path, root)
                    yield path, domain, entry.get("url"), kind


def discover_additional_files(
    root: Path, seen: set[Path], include_images: bool = True
) -> Iterable[tuple[Path, Optional[str], Optional[str], Optional[str]]]:
    """Find docs/images under scraped_data not referenced in metadata."""
    allowed_exts = (
        SUPPORTED_TEXT_EXTS
        | SUPPORTED_HTML_EXTS
        | SUPPORTED_DOC_EXTS
        | SUPPORTED_PDF_EXTS
        | (SUPPORTED_IMAGE_EXTS if include_images else set())
    )
    for path in sorted(root.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in allowed_exts:
            continue
        resolved = path.resolve()
        if resolved in seen:
            continue
        yield path, None, None, None


def extract_path(
    path: Path, ocr_pdf_pages: int, default_kind: Optional[str]
) -> Optional[tuple[str, str]]:
    ext = path.suffix.lower()
    try:
        if ext in SUPPORTED_TEXT_EXTS:
            return extract_text_file(path), default_kind or "text"
        if ext in SUPPORTED_HTML_EXTS:
            return extract_html_file(path), default_kind or "html"
        if ext in SUPPORTED_DOC_EXTS:
            return extract_docx_file(path), default_kind or "docx"
        if ext in SUPPORTED_PDF_EXTS:
            return extract_pdf_file(path, ocr_page_limit=ocr_pdf_pages), default_kind or "pdf"
        if ext in SUPPORTED_IMAGE_EXTS:
            return extract_image_file(path), default_kind or "image_ocr"
    except Exception as exc:  # noqa: BLE001
        logging.warning("Failed to extract %s: %s", path, exc)
    return None


def iter_sources(
    root: Path, ocr_pdf_pages: int, include_images: bool
) -> Iterable[ExtractedDocument]:
    seen: set[Path] = set()

    for path, domain, url, kind in load_metadata_entries(root):
        if not path.exists():
            logging.debug("Skipping missing path from metadata: %s", path)
            continue

        resolved = path.resolve()
        if resolved in seen:
            continue
        seen.add(resolved)

        extracted = extract_path(path, ocr_pdf_pages, default_kind=kind)
        if not extracted:
            continue
        text, doc_kind = extracted
        if not text:
            continue
        language = detect_language(text)
        yield ExtractedDocument(
            source_path=str(path.relative_to(root)),
            source_domain=domain,
            source_url=url,
            kind=doc_kind,
            language=language,
            text=text,
            length=len(text),
        )

    for path, domain, url, kind in discover_additional_files(
        root, seen=seen, include_images=include_images
    ):
        extracted = extract_path(path, ocr_pdf_pages, default_kind=kind)
        if not extracted:
            continue
        text, doc_kind = extracted
        if not text:
            continue

        relative = path.relative_to(root)
        domain = domain or (relative.parts[0] if relative.parts else "")
        language = detect_language(text)
        yield ExtractedDocument(
            source_path=str(relative),
            source_domain=domain,
            source_url=url,
            kind=doc_kind,
            language=language,
            text=text,
            length=len(text),
        )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Extract raw text from scraped_data (html, pdfs, docx, images)."
    )
    parser.add_argument(
        "--root",
        type=Path,
        default=Path("scraped_data"),
        help="Root directory containing scraped data.",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("processed_data/extracted_documents.jsonl"),
        help="Where to write the extracted JSONL file.",
    )
    parser.add_argument(
        "--limit",
        type=int,
        default=None,
        help="Process only this many files (useful for dry runs).",
    )
    parser.add_argument(
        "--ocr-pdf-pages",
        type=int,
        default=2,
        help="How many pages to OCR when PDFs have no extractable text.",
    )
    parser.add_argument(
        "--no-images",
        action="store_true",
        help="Skip OCR on images and image-based PDFs.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

    args.output.parent.mkdir(parents=True, exist_ok=True)
    include_images = not args.no_images

    written = 0
    with args.output.open("w", encoding="utf-8") as f:
        for idx, doc in enumerate(
            tqdm(
                iter_sources(args.root, args.ocr_pdf_pages, include_images),
                desc="Extracting documents",
            )
        ):
            if args.limit is not None and idx >= args.limit:
                break
            f.write(json.dumps(asdict(doc), ensure_ascii=False) + "\n")
            written += 1

    logging.info(
        "Saved %d extracted documents to %s (limit=%s, include_images=%s)",
        written,
        args.output,
        args.limit,
        include_images,
    )


if __name__ == "__main__":
    main()
