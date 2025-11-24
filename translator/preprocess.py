from __future__ import annotations

import argparse
import json
import logging
import os
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Callable, Dict, Iterable, Optional

import pdfplumber
import pytesseract
from docx import Document
from dotenv import load_dotenv
from langdetect import DetectorFactory, LangDetectException, detect
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_groq import ChatGroq
from langchain_openai import ChatOpenAI
from PIL import Image
from tqdm import tqdm

# langdetect can be non-deterministic unless we set the seed.
DetectorFactory.seed = 0
load_dotenv()


SUPPORTED_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif"}
SUPPORTED_TEXT_EXTS = {".txt", ".html", ".htm"}
SUPPORTED_DOC_EXTS = {".docx"}
SUPPORTED_PDF_EXTS = {".pdf"}


@dataclass
class ProcessedDocument:
    source_path: str
    source_domain: str
    file_type: str
    language: str
    translated: bool
    translation_source: Optional[str]
    text: str
    original_length: int
    translated_length: int


def clean_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    cleaned = "\n".join(line for line in lines if line)
    return cleaned.strip()


def extract_text_file(path: Path) -> str:
    return clean_text(path.read_text(encoding="utf-8", errors="ignore"))


def extract_docx_file(path: Path) -> str:
    document = Document(path)
    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    return clean_text("\n".join(parts))


def ocr_image(img: Image.Image) -> str:
    return pytesseract.image_to_string(img)


def extract_image_file(path: Path) -> str:
    with Image.open(path) as img:
        return clean_text(ocr_image(img))


def extract_pdf_file(path: Path, ocr_page_limit: int = 2) -> str:
    text_parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    full_text = clean_text("\n".join(text_parts))
    if full_text:
        return full_text

    # Fallback to OCR for scanned PDFs.
    ocr_parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages[:ocr_page_limit]:
            page_image = page.to_image(resolution=300)
            ocr_parts.append(ocr_image(page_image.original))
    return clean_text("\n".join(ocr_parts))


def detect_language(text: str) -> str:
    try:
        return detect(text)
    except LangDetectException:
        return "unknown"


def chunk_text(text: str, max_chars: int = 1200) -> list[str]:
    chunks: list[str] = []
    current: list[str] = []
    current_len = 0
    for paragraph in text.split("\n"):
        if not paragraph.strip():
            continue
        if current_len + len(paragraph) > max_chars and current:
            chunks.append("\n".join(current))
            current = []
            current_len = 0
        current.append(paragraph)
        current_len += len(paragraph)
    if current:
        chunks.append("\n".join(current))
    return chunks


class NepaliTranslator:
    def __init__(
        self,
        provider: str = "none",
        model: Optional[str] = None,
        temperature: float = 0.0,
    ) -> None:
        self.provider = provider
        self.translation_source = None
        self.llm = self._build_llm(provider, model, temperature)

    def _build_llm(self, provider: str, model: Optional[str], temperature: float):
        if provider == "groq":
            api_key = os.getenv("GROQ_API_KEY")
            if not api_key:
                logging.warning("GROQ_API_KEY not found; translation disabled.")
                return None
            self.translation_source = f"groq:{model or 'llama-3.3-70b-versatile'}"
            return ChatGroq(model=model or "llama-3.3-70b-versatile", temperature=temperature)

        if provider == "openai":
            api_key = os.getenv("OPENAI_API_KEY")
            if not api_key:
                logging.warning("OPENAI_API_KEY not found; translation disabled.")
                return None
            self.translation_source = f"openai:{model or 'gpt-4o-mini'}"
            return ChatOpenAI(model=model or "gpt-4o-mini", temperature=temperature)

        logging.info("Translation provider set to none; skipping translation.")
        return None

    def translate(self, text: str) -> tuple[str, bool]:
        if not self.llm:
            return text, False

        prompt = [
            SystemMessage(
                content=(
                    "You are a translation assistant. Translate Nepali text to clear, concise English. "
                    "Return only the translation without additional commentary."
                )
            )
        ]

        translated_chunks: list[str] = []
        for chunk in chunk_text(text):
            response = self.llm.invoke(prompt + [HumanMessage(content=chunk)])
            translated_chunks.append(response.content.strip())
        return "\n".join(translated_chunks), True


def build_extractors(ocr_page_limit: int) -> Dict[str, Callable[[Path], str]]:
    return {
        ext: extract_image_file for ext in SUPPORTED_IMAGE_EXTS
    } | {
        ext: extract_text_file for ext in SUPPORTED_TEXT_EXTS
    } | {
        ext: extract_docx_file for ext in SUPPORTED_DOC_EXTS
    } | {
        ext: (lambda path, limit=ocr_page_limit: extract_pdf_file(path, limit))
        for ext in SUPPORTED_PDF_EXTS
    }


def iter_source_files(
    root: Path, allowed_exts: set[str], limit: Optional[int] = None
) -> Iterable[Path]:
    """Yield only supported files in a stable order, respecting the optional limit."""
    files = sorted(
        p for p in root.rglob("*") if p.is_file() and p.suffix.lower() in allowed_exts
    )
    for idx, path in enumerate(files):
        if limit and idx >= limit:
            break
        yield idx, path


def process_file(
    path: Path,
    extractors: Dict[str, Callable[[Path], str]],
    translator: NepaliTranslator,
    root: Path,
) -> Optional[ProcessedDocument]:
    ext = path.suffix.lower()
    extractor = extractors.get(ext)
    if not extractor:
        return None

    text = extractor(path)
    if not text:
        return None

    language = detect_language(text)
    translated_text, translated = (
        translator.translate(text) if language == "ne" else (text, False)
    )

    relative_path = path.relative_to(root)
    source_domain = relative_path.parts[0] if relative_path.parts else ""

    return ProcessedDocument(
        source_path=str(path.relative_to(root)),
        source_domain=source_domain,
        file_type=ext,
        language=language,
        translated=translated,
        translation_source=translator.translation_source if translated else None,
        text=translated_text,
        original_length=len(text),
        translated_length=len(translated_text),
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Extract, translate (Nepali to English), and normalize scraped data."
    )
    parser.add_argument(
        "--root",
        type=Path,
        default=Path("scraped_data"),
        help="Root directory containing scraped data.",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("processed_data/english_documents.jsonl"),
        help="Where to write the processed JSONL file.",
    )
    parser.add_argument(
        "--translator",
        choices=["groq", "openai", "none"],
        default="none",
        help="Translation backend to use.",
    )
    parser.add_argument(
        "--translator-model",
        type=str,
        default=None,
        help="Optional model name override for the translator.",
    )
    parser.add_argument(
        "--limit",
        type=int,
        default=None,
        help="Process only this many files (useful for dry runs).",
    )
    parser.add_argument(
        "--start-index",
        type=int,
        default=0,
        help="Skip files before this index to resume a previous run.",
    )
    parser.add_argument(
        "--append",
        action="store_true",
        help="Append to the output file instead of overwriting (useful when resuming).",
    )
    parser.add_argument(
        "--ocr-pdf-pages",
        type=int,
        default=2,
        help="How many pages to OCR when PDFs have no extractable text.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

    extractors = build_extractors(args.ocr_pdf_pages)
    allowed_exts = set(extractors.keys())
    translator = NepaliTranslator(provider=args.translator, model=args.translator_model)

    root = args.root
    args.output.parent.mkdir(parents=True, exist_ok=True)

    processed = 0
    mode = "a" if args.append else "w"
    with args.output.open(mode, encoding="utf-8") as f:
        for idx, path in tqdm(
            iter_source_files(root, allowed_exts, args.limit), desc="Processing files"
        ):
            if idx < args.start_index:
                continue
            doc = process_file(path, extractors, translator, root=root)
            if doc:
                f.write(json.dumps(asdict(doc), ensure_ascii=False) + "\n")
                processed += 1

    logging.info(
        "Saved %d documents to %s (start_index=%d, limit=%s, append=%s)",
        processed,
        args.output,
        args.start_index,
        args.limit,
        args.append,
    )


if __name__ == "__main__":
    main()
